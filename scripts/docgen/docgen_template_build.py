"""
Build or modify a .docx document template programmatically.

Supports:
  - Creating a new .docx from a JSON layout spec (tables, paragraphs, images)
  - Replacing tokens in an existing .docx (useful for branding updates)
  - Adding static images to an existing .docx
  - Listing/auditing structure of an existing .docx

Requires: pip install python-docx

Usage:
  python scripts/docgen/docgen_template_build.py create layout.json --output template.docx
  python scripts/docgen/docgen_template_build.py replace template.docx --tokens '{"OldToken": "NewToken"}'
  python scripts/docgen/docgen_template_build.py audit template.docx
  python scripts/docgen/docgen_template_build.py --example > layout.json
"""
import argparse
import json
import re
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print(
        "ERROR: python-docx required. Install with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


LAYOUT_EXAMPLE = {
    "_comment": "Layout spec for docgen_template_build.py create",
    "page": {"margin_inches": 0.75},
    "elements": [
        {
            "type": "image",
            "path": "/path/to/logo.png",
            "width_inches": 2.5,
            "alignment": "left",
        },
        {"type": "spacer"},
        {
            "type": "heading",
            "text": "INVOICE",
            "level": 1,
            "alignment": "right",
            "bold": True,
        },
        {"type": "spacer"},
        {
            "type": "table",
            "columns": 2,
            "style": "Table Grid",
            "rows": [
                ["Invoice Number:", "{{InvoiceNumber}}"],
                ["Invoice Date:", "{{InvoiceDate}}"],
                ["Due Date:", "{{DueDate}}"],
            ],
        },
        {"type": "spacer"},
        {
            "type": "paragraph",
            "text": "Bill To: {{AccountName}}",
            "bold": False,
            "size_pt": 11,
        },
        {"type": "spacer"},
        {
            "_comment": "Repeating section for line items",
            "type": "table",
            "columns": 4,
            "header": ["Description", "Quantity", "Unit Price", "Amount"],
            "rows": [
                [
                    "{{#InvoiceLines}}",
                    "",
                    "",
                    "",
                ],
                [
                    "{{Description}}",
                    "{{Quantity}}",
                    "{{UnitPrice}}",
                    "{{TotalAmount}}",
                ],
                ["{{/InvoiceLines}}", "", "", ""],
            ],
        },
        {"type": "spacer"},
        {
            "type": "paragraph",
            "text": "Total: {{TotalAmount}}",
            "bold": True,
            "size_pt": 12,
            "alignment": "right",
        },
    ],
}


def apply_alignment(paragraph, alignment_str):
    alignments = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    if alignment_str and alignment_str.lower() in alignments:
        paragraph.alignment = alignments[alignment_str.lower()]


def create_from_layout(layout, output_path):
    doc = Document()

    page = layout.get("page", {})
    margin = page.get("margin_inches", 1.0)
    for section in doc.sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

    for elem in layout.get("elements", []):
        elem_type = elem.get("type", "")

        if elem_type == "spacer":
            doc.add_paragraph("")

        elif elem_type == "paragraph":
            p = doc.add_paragraph()
            run = p.add_run(elem.get("text", ""))
            if elem.get("bold"):
                run.bold = True
            if elem.get("italic"):
                run.italic = True
            if elem.get("size_pt"):
                run.font.size = Pt(elem["size_pt"])
            if elem.get("color"):
                rgb = elem["color"].lstrip("#")
                run.font.color.rgb = RGBColor(
                    int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16)
                )
            apply_alignment(p, elem.get("alignment"))

        elif elem_type == "heading":
            p = doc.add_heading(elem.get("text", ""), level=elem.get("level", 1))
            if elem.get("alignment"):
                apply_alignment(p, elem["alignment"])

        elif elem_type == "image":
            p = doc.add_paragraph()
            run = p.add_run()
            try:
                width = Inches(elem["width_inches"]) if elem.get("width_inches") else None
                run.add_picture(elem["path"], width=width)
            except FileNotFoundError:
                run.add_text(f"[IMAGE NOT FOUND: {elem['path']}]")
            apply_alignment(p, elem.get("alignment"))

        elif elem_type == "table":
            cols = elem.get("columns", 2)
            rows_data = elem.get("rows", [])
            header = elem.get("header")

            total_rows = len(rows_data) + (1 if header else 0)
            table = doc.add_table(rows=total_rows, cols=cols)
            table.style = elem.get("style", "Table Grid")

            row_offset = 0
            if header:
                for i, cell_text in enumerate(header[:cols]):
                    cell = table.rows[0].cells[i]
                    cell.text = cell_text
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                row_offset = 1

            for r, row_data in enumerate(rows_data):
                for c, cell_text in enumerate(row_data[:cols]):
                    table.rows[r + row_offset].cells[c].text = cell_text

        elif elem_type == "page_break":
            doc.add_page_break()

    doc.save(output_path)
    print(f"Created: {output_path}")


def replace_tokens(docx_path, token_map, output_path):
    doc = Document(docx_path)
    count = 0

    for para in doc.paragraphs:
        full_text = para.text
        for old, new in token_map.items():
            if old in full_text:
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in token_map.items():
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                count += 1

    doc.save(output_path or docx_path)
    print(f"Replaced {count} occurrence(s) in: {output_path or docx_path}")


def audit_docx(docx_path):
    doc = Document(docx_path)
    print(f"Auditing: {docx_path}")
    print(f"  Sections: {len(doc.sections)}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")

    token_re = re.compile(r"\{\{([^}]+)\}\}")
    tokens_found = set()
    images_found = 0

    for para in doc.paragraphs:
        for match in token_re.finditer(para.text):
            tokens_found.add(match.group(1))
        for run in para.runs:
            if run.element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            ) or run.element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
            ):
                images_found += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for match in token_re.finditer(para.text):
                        tokens_found.add(match.group(1))

    print(f"  Embedded images: {images_found}")
    print(f"  Unique tokens: {len(tokens_found)}")
    if tokens_found:
        fields = sorted(t for t in tokens_found if not t.startswith(("#", "/", "IMG_")))
        sections = sorted(t[1:] for t in tokens_found if t.startswith("#"))
        images = sorted(t for t in tokens_found if t.startswith("IMG_"))
        if fields:
            print(f"    Fields: {fields}")
        if sections:
            print(f"    Sections: {sections}")
        if images:
            print(f"    Images: {images}")


def main():
    parser = argparse.ArgumentParser(
        description="Build or modify .docx document templates"
    )
    subparsers = parser.add_subparsers(dest="command")

    create_p = subparsers.add_parser("create", help="Create .docx from layout spec")
    create_p.add_argument("layout", help="Path to layout JSON spec")
    create_p.add_argument("--output", "-o", required=True, help="Output .docx path")

    replace_p = subparsers.add_parser(
        "replace",
        help="Replace tokens in existing .docx (single-run tokens only)",
        epilog="Limitation: only replaces tokens contained within a single Word "
               "run. If Word splits a token across multiple runs (e.g. due to "
               "spell-check or formatting), the replacement will not match. "
               "Headers and footers are not searched. Use 'audit' to verify "
               "which tokens are visible.",
    )
    replace_p.add_argument("docx", help="Path to .docx file")
    replace_p.add_argument(
        "--tokens", required=True, help='JSON map: {"old": "new", ...}'
    )
    replace_p.add_argument("--output", "-o", help="Output path (default: overwrite)")

    audit_p = subparsers.add_parser("audit", help="Audit .docx structure")
    audit_p.add_argument("docx", help="Path to .docx file")

    parser.add_argument(
        "--example", action="store_true", help="Print example layout spec and exit"
    )

    args = parser.parse_args()

    if args.example:
        print(json.dumps(LAYOUT_EXAMPLE, indent=2))
        return

    if not args.command:
        parser.print_help()
        return

    if args.command == "create":
        try:
            with open(args.layout) as f:
                layout = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"ERROR: {e}", file=sys.stderr)
            sys.exit(1)
        create_from_layout(layout, args.output)

    elif args.command == "replace":
        try:
            token_map = json.loads(args.tokens)
        except json.JSONDecodeError as e:
            print(f"ERROR: Invalid JSON in --tokens: {e}", file=sys.stderr)
            sys.exit(1)
        replace_tokens(args.docx, token_map, args.output)

    elif args.command == "audit":
        audit_docx(args.docx)


if __name__ == "__main__":
    main()
